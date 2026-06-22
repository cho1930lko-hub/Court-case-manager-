"""DOCX generation for the 5 court document types using python-docx."""
import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def _set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '333333')
        borders.append(el)
    # Per CT_TblPrBase sequence, tblBorders must come after tblStyle, tblpPr,
    # tblOverlap, bidiVisual, tblStyleRowBandSize, tblStyleColBandSize, tblW,
    # jc, tblCellSpacing, tblInd — and before shd, tblLayout, tblCellMar, tblLook.
    preceding_tags = ['tblStyle', 'tblpPr', 'tblOverlap', 'bidiVisual',
                       'tblStyleRowBandSize', 'tblStyleColBandSize', 'tblW',
                       'jc', 'tblCellSpacing', 'tblInd']
    insert_after = None
    for tag in preceding_tags:
        found = tbl_pr.find(qn(f'w:{tag}'))
        if found is not None:
            insert_after = found
    if insert_after is not None:
        insert_after.addnext(borders)
    else:
        tbl_pr.insert(0, borders)


def _base_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Nirmala UI'
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), 'Nirmala UI')
    rfonts.set(qn('w:cs'), 'Nirmala UI')

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    return doc


def _add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.size = Pt(16)
    return p


def _add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(13)
    return p


def _add_info_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(6)
    table.columns[1].width = Cm(11)
    _set_table_borders(table)
    for label, value in rows:
        row = table.add_row()
        c0, c1 = row.cells
        c0.width = Cm(6)
        c1.width = Cm(11)
        _set_cell_shading(c0, "F2F2F2")
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.size = Pt(11)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(value or "")
        r1.font.size = Pt(11)
    return table


def _add_section_label(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p


def _add_name_pata(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(12)
    return p


def _add_body_text(doc, text, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(11.5)
    return p


def _add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.underline = True
    run.font.size = Pt(10.5)
    return p


def _add_sign_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(line)
        run.font.size = Pt(11.5)


def _common_table_rows(f):
    return [
        ("राज्य बनाम", f.get("rajya_banam", "")),
        ("मु0न0", f.get("mu_no", "")),
        ("मु0अ0सं0", f.get("mu_apradh_sankhya", "")),
        ("धारा", f.get("dhara", "")),
        ("थाना", f.get("thana", "")),
        ("जनपद", f.get("janpad", "")),
        ("तारीख पेशी", f.get("tareekh_peshi", "")),
    ]


def build_nbw(f) -> Document:
    doc = _base_doc()
    _add_title(doc, "गैर जमानती वारंट (NBW) साक्षी")
    _add_subtitle(doc, f"न्यायालय : {f.get('nyayalaya_pad','')} {f.get('janpad','')}")
    _add_info_table(doc, _common_table_rows(f))
    _add_section_label(doc, "नाम पता साक्षी:")
    _add_name_pata(doc, f.get("naam_pata", ""))
    _add_body_text(doc, "उपरोक्त साक्षी को नियत तारीख पेशी से पूर्व या तक गिरफ्तार कर मेरे समक्ष पेश करें।")
    _add_note(doc, "नोट:- तामीलकर्ता साक्षी का मोबाइल नंबर अवश्य अंकित करें।")
    _add_sign_block(doc, [f.get('nyayalaya_pad',''), f.get('janpad',''), f.get('aaj_ki_tareekh','')])
    return doc


def build_bw(f) -> Document:
    doc = _base_doc()
    _add_title(doc, "जमानती वारंट (BW) साक्षी")
    _add_subtitle(doc, f"न्यायालय : {f.get('nyayalaya_pad','')} {f.get('janpad','')}")
    _add_info_table(doc, _common_table_rows(f))
    _add_section_label(doc, "नाम पता साक्षी:")
    _add_name_pata(doc, f.get("naam_pata", ""))
    _add_body_text(doc, f"उपरोक्त साक्षी से {f.get('jamanat_rashi','')} की दो जमानतें तथा इतनी ही धनराशि का बंधपत्र न्यायालय उपस्थित होने के बाबत दाखिल कराएं। ऐसा ना करने पर उपर्युक्त साक्षी को गिरफ्तार कर मेरे समक्ष पेश करें।")
    _add_note(doc, "नोट:- तामीलकर्ता साक्षी का मोबाइल नंबर अवश्य अंकित करें।")
    _add_sign_block(doc, [f.get('nyayalaya_pad',''), f.get('janpad',''), f.get('aaj_ki_tareekh','')])
    return doc


def build_mrityu(f) -> Document:
    doc = _base_doc()
    _add_title(doc, "मृत्यु आख्या")
    _add_subtitle(doc, f"न्यायालय : {f.get('nyayalaya_pad','')} {f.get('janpad','')}")
    _add_info_table(doc, _common_table_rows(f))
    _add_section_label(doc, "नाम पता अभियुक्त:")
    _add_name_pata(doc, f.get("naam_pata", ""))
    _add_body_text(doc, "एतद् द्वारा आदेशित किया जाता है कि उपर्युक्त अभियुक्त की मृत्यु से संबंधित प्रमाण-पत्र नियत तिथि पर न्यायालय के समक्ष अनिवार्य रूप से प्रस्तुत किया जाए।")
    _add_sign_block(doc, [f.get('nyayalaya_pad',''), f.get('janpad',''), f.get('aaj_ki_tareekh','')])
    return doc


def build_saman(f) -> Document:
    doc = _base_doc()
    _add_title(doc, "समन साक्षी")
    _add_subtitle(doc, f"न्यायालय : {f.get('nyayalaya_pad','')} {f.get('janpad','')}")
    _add_info_table(doc, _common_table_rows(f))
    _add_section_label(doc, "नाम पता साक्षी:")
    _add_name_pata(doc, f.get("naam_pata", ""))
    _add_body_text(doc, f"साक्षी मुकदमा उपरोक्त के संबंध में अपना साक्ष्य देने के लिए अपने पहचान पत्र / आधार कार्ड व एक पासपोर्ट साइज फोटो के साथ उपर्युक्त नियत तारीख पेशी पर {f.get('peshi_samay','')} उपस्थित होना सुनिश्चित करें। ऐसा करने में कोई त्रुटि न हो।")
    _add_note(doc, "नोट:- तामीलकर्ता साक्षी का मोबाइल नंबर अवश्य अंकित करें।")
    _add_sign_block(doc, [f.get('nyayalaya_pad',''), f.get('janpad',''), f.get('aaj_ki_tareekh','')])
    return doc


def build_vc(f) -> Document:
    doc = _base_doc()

    p = doc.add_paragraph()
    p.add_run("प्रेषक,").font.size = Pt(11.5)
    p2 = doc.add_paragraph()
    p2.add_run(f"\t{f.get('praapak_pad','')},").font.size = Pt(11.5)
    p3 = doc.add_paragraph()
    p3.add_run(f"\t{f.get('janpad','')} ।").font.size = Pt(11.5)

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.add_run("सेवा में,").font.size = Pt(11.5)
    p5 = doc.add_paragraph()
    p5.add_run(f"\t{f.get('seva_me_pad','')},").font.size = Pt(11.5)
    p6 = doc.add_paragraph()
    p6.add_run(f"\tजनपद {f.get('seva_me_janpad','')}।").font.size = Pt(11.5)

    doc.add_paragraph()
    p7 = doc.add_paragraph()
    r7 = p7.add_run(f"विषय - {f.get('vishay','')}")
    r7.bold = True
    r7.font.size = Pt(11.5)

    doc.add_paragraph()
    _add_body_text(doc, "महोदय,")
    _add_body_text(doc, "सविनय निवेदन है कि इस न्यायालय में लंबित निम्नलिखित विवरण के अनुसार साक्षी का साक्ष्य अंकन जरिए वीडियो कांफ्रेंसिंग कराया जाना है:")

    _add_info_table(doc, [
        ("सत्र परीक्षण संख्या", f.get("satra_pareeksha_sankhya", "")),
        ("मुकदमा अपराध संख्या", f.get("mu_apradh_sankhya", "")),
        ("राज्य बनाम", f.get("rajya_banam", "")),
        ("धारा", f.get("dhara", "")),
        ("थाना", f.get("thana", "")),
        ("जनपद", f.get("janpad", "")),
        ("तारीख पेशी", f.get("tareekh_peshi", "")),
    ])

    _add_body_text(doc, f"अतः आपसे निवेदन है कि साक्षी का साक्ष्य अंकन {f.get('vc_tareekh','')} को जिला न्यायालय {f.get('vc_sthan','')} में वीडियो कांफ्रेंसिंग के माध्यम से करवाने हेतु आवश्यक निर्देश प्रदान करने की कृपा करें।")

    _add_sign_block(doc, ["भवदीय,", f.get('praapak_pad',''), f.get('janpad',''), f.get('aaj_ki_tareekh','')])

    doc.add_paragraph()
    p8 = doc.add_paragraph()
    r8 = p8.add_run("प्रतिलिपि-")
    r8.bold = True
    r8.font.size = Pt(11)

    _add_body_text(doc, f"1. {f.get('prati_pad','')} जनपद {f.get('prati_janpad','')} को इस आशय के साथ प्रेषित कि संबंधित अधीनस्थ को उपरोक्तानुसार निर्देश प्रदान करने का कष्ट करें।")
    _add_body_text(doc, f"2. {f.get('naam_pata','')} दिनांक {f.get('vc_tareekh','')} को समय {f.get('vc_samay','')} पर जिला न्यायालय {f.get('vc_sthan','')} में अपनी पहचान संबंधित दस्तावेज के साथ वीडियो कांफ्रेंसिंग हेतु उपस्थित रहें।")

    _add_sign_block(doc, [f.get('praapak_pad',''), f.get('janpad',''), f.get('aaj_ki_tareekh','')])

    return doc


BUILDERS = {
    "nbw": build_nbw,
    "bw": build_bw,
    "mrityu": build_mrityu,
    "saman": build_saman,
    "vc": build_vc,
}


def generate_docx_bytes(doc_type: str, fields: dict) -> bytes:
    builder = BUILDERS.get(doc_type)
    if not builder:
        raise ValueError(f"Unknown doc type: {doc_type}")
    doc = builder(fields)

    # Fix w:zoom missing percent attribute (python-docx default omits it,
    # which some strict validators flag even though Word/LibreOffice accept it).
    settings = doc.settings.element
    zoom = settings.find(qn('w:zoom'))
    if zoom is not None and zoom.get(qn('w:percent')) is None:
        zoom.set(qn('w:percent'), '100')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

