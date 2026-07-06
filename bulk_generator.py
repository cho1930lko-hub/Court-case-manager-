"""
Bulk DOCX Generator — आपकी sheet के actual columns के हिसाब से:
ST NO | मु0अ0स0 | धारा | बनाम | न्यायालय | तलब साक्षी | अगली तारीख पेशी

तलब साक्षी column में prefix:
  BW xxxxx   → जमानती वारंट
  NBW xxxxx  → गैर जमानती वारंट
  MRITYU xxx → मृत्यु आख्या
  बिना prefix → साधारण समन साक्षी

VC Sheet अलग sheet में (optional)
"""

import re
import io
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


# ── Column mapping (आपकी sheet → internal keys) ─────────────────────────────
COL = {
    "st_no":    "ST NO",
    "mu_apradh":"मु0अ0स0",
    "dhara":    "धारा",
    "banam":    "बनाम",
    "nyayalaya":"न्यायालय",
    "saakshi":  "तलब साक्षी",
    "tareekh":  "अगली तारीख पेशी",
}

REQUIRED_COLS = list(COL.values())

# न्यायालय keywords → display name
NYAYALAYA_MAP = {
    "ASJ":   ("अपर सत्र न्यायाधीश", True),
    "DJ":    ("जिला एवं सत्र न्यायाधीश", True),
    "POCSO": ("विशेष न्यायालय POCSO", True),
    "CJM":   ("मुख्य न्यायिक मजिस्ट्रेट", False),
    "JM":    ("न्यायिक मजिस्ट्रेट", False),
    "ACJM":  ("अपर मुख्य न्यायिक मजिस्ट्रेट", False),
    "MM":    ("महानगर मजिस्ट्रेट", False),
}


def resolve_court(raw, janpad):
    raw_up = str(raw).upper().strip()
    for key, (name, is_session) in NYAYALAYA_MAP.items():
        if key in raw_up:
            color = RGBColor(0, 0, 139) if is_session else None
            return f"{name} जनपद {janpad}", color
    return f"{raw} जनपद {janpad}", None


def clean_col(name):
    return re.sub(r'\s+', ' ', str(name)).strip()


def detect_type(field):
    w = str(field).strip()
    if re.match(r'^NBW', w, re.IGNORECASE): return "NBW"
    if re.match(r'^BW',  w, re.IGNORECASE): return "BW"
    if re.match(r'^MRITYU', w, re.IGNORECASE): return "MRITYU"
    return "SAMAN"


def clean_prefix(name):
    return re.sub(r'^\s*(BW|NBW|MRITYU|VC)\s*', '', name, flags=re.IGNORECASE).strip()


def fmt_date(val):
    if isinstance(val, datetime): return val.strftime('%d/%m/%Y')
    s = str(val).strip()
    for fmt in ('%d/%m/%Y', '%d-%m-%Y', '%Y-%m-%d'):
        try: return datetime.strptime(s, fmt).strftime('%d/%m/%Y')
        except: pass
    return s


# ── Document helpers ─────────────────────────────────────────────────────────

def fmt(doc, text, bold=False, underline=False, italic=False,
        size=17, align=None, color=None, sb=0, sa=2):
    p = doc.add_paragraph()
    r = p.add_run(str(text))
    r.font.name = "Kokila"
    r.font.size = Pt(size)
    r.bold = bold; r.underline = underline; r.italic = italic
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    return p


def set_cell(cell, text, bold=False, size=17):
    cell.text = str(text)
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = "Kokila"
            run.font.size = Pt(size)
            run.bold = bold


def build_table(doc, st_no, mu_no, banam, dhara, janpad, tareekh):
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    rows_data = [
        ("राज्य बनाम", banam),
        ("ST NO / मु0न0", st_no),
        ("मु0अ0स0", mu_no),
        ("धारा", dhara),
        ("जनपद", janpad),
        ("तारीख पेशी", tareekh),
    ]
    for i, (f, v) in enumerate(rows_data):
        set_cell(table.cell(i, 0), f)
        set_cell(table.cell(i, 1), v)


def base_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width   = Cm(29.7)
    sec.page_height  = Cm(21.0)
    sec.left_margin  = Cm(1.27)
    sec.right_margin = Cm(1.27)
    sec.top_margin   = Cm(1.27)
    sec.bottom_margin= Cm(1.27)
    return doc


# ── Per-summon block renderer ────────────────────────────────────────────────

TITLES = {
    "SAMAN":  "समन साक्षी",
    "BW":     "जमानती वारंट (BW) साक्षी",
    "NBW":    "गैर जमानती वारंट (NBW) साक्षी",
    "MRITYU": "मृत्यु आख्या",
}

BODY = {
    "SAMAN": ("साक्षी मुकदमा उपरोक्त के संबंध में अपना साक्ष्य देने के लिए अपने पहचान पत्र / "
              "आधार कार्ड व एक पासपोर्ट साइज फोटो के साथ उपर्युक्त नियत तारीख पेशी पर "
              "सुबह 10:30 बजे उपस्थित होना सुनिश्चित करें। ऐसा करने में कोई त्रुटि न हो।"),
    "BW":    ("उपरोक्त साक्षी से 5-5 हजार की दो जमानतें तथा इतनी ही धनराशि का बंधपत्र "
              "न्यायालय में साक्ष्य हेतु उपस्थित होने के बाबत दाखिल कराएं। ऐसा ना करने पर "
              "उपर्युक्त साक्षी को गिरफ्तार कर नियत पेशी पर इस न्यायालय पर प्रस्तुत करें।"),
    "NBW":   "उपरोक्त साक्षी को नियत तारीख पेशी से पूर्व या तक गिरफ्तार कर मेरे समक्ष पेश करें।",
    "MRITYU":("एतद् द्वारा आदेशित किया जाता है कि उपर्युक्त अभियुक्त की मृत्यु से संबंधित "
              "प्रमाण-पत्र नियत तिथि पर न्यायालय के समक्ष अनिवार्य रूप से प्रस्तुत किया जाए।"),
}


def render_block(doc, row_data, stype, janpad):
    court_disp, court_color = resolve_court(row_data[COL["nyayalaya"]], janpad)
    st_no   = str(row_data.get(COL["st_no"],    "")).strip()
    mu_no   = str(row_data.get(COL["mu_apradh"],"")).strip()
    banam   = str(row_data.get(COL["banam"],    "")).strip()
    dhara   = str(row_data.get(COL["dhara"],    "")).strip()
    tareekh = fmt_date(row_data.get(COL["tareekh"], ""))
    witness = clean_prefix(str(row_data.get(COL["saakshi"], "")).strip())

    fmt(doc, TITLES[stype], bold=True, underline=True, size=22,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER, sa=2)

    p = fmt(doc, f"न्यायालय : {court_disp}", bold=True, size=18,
            align=WD_PARAGRAPH_ALIGNMENT.CENTER, sa=4)
    if court_color:
        p.runs[0].font.color.rgb = court_color

    build_table(doc, st_no, mu_no, banam, dhara, janpad, tareekh)

    label = "नाम पता अभियुक्त:" if stype == "MRITYU" else "नाम पता साक्षी:"
    fmt(doc, "\n" + label, bold=True, size=17, sa=1)
    fmt(doc, witness, italic=True, size=17, sa=2)
    fmt(doc, BODY[stype], size=16, sa=2)

    if stype != "MRITYU":
        fmt(doc, "नोट:- तामीलकर्ता साक्षी का मोबाइल नंबर अवश्य अंकित करें।",
            bold=True, underline=True, size=15, sa=2)

    today = datetime.today().strftime('%d/%m/%Y')
    fmt(doc, f"\n{court_disp}\n{today}",
        align=WD_PARAGRAPH_ALIGNMENT.RIGHT, size=17)


# ── VC Letter renderer ───────────────────────────────────────────────────────

def render_vc(doc, row, janpad):
    court_raw   = str(row.get("न्यायालय", "जिला एवं सत्र")).strip()
    st_no       = str(row.get("ST NO", "")).strip()
    mu_no       = str(row.get("मु0अ0स0", "")).strip()
    banam       = str(row.get("बनाम", "")).strip()
    dhara       = str(row.get("धारा", "")).strip()
    tareekh     = fmt_date(row.get("अगली तारीख पेशी", ""))
    witness     = str(row.get("तलब साक्षी", "")).strip()
    seva_janpad = str(row.get("सेवा जनपद", "")).strip()
    vc_sthan    = str(row.get("VC स्थान", "पुलिस कार्यालय")).strip()
    vc_samay    = str(row.get("VC समय", "12:30 PM")).strip()
    prati_pad   = str(row.get("प्रति पद", "पुलिस अधीक्षक")).strip()
    today       = datetime.today().strftime('%d/%m/%Y')

    praapak, _ = resolve_court(court_raw, janpad)

    fmt(doc, f"प्रेषक,\n    {praapak} ।", size=16, sa=2)
    fmt(doc, f"\nसेवा में,\n    जिला एवं सत्र न्यायाधीश,\n    जनपद {seva_janpad or '___'}।", size=16, sa=2)
    fmt(doc, f"\nविषय - {witness} का साक्ष्य जरिए वीडियो कांफ्रेंसिंग अंकित कराने के सम्बन्ध में।",
        bold=True, size=16, sa=4)
    fmt(doc, "महोदय,", size=16, sa=1)
    fmt(doc, ("सविनय निवेदन है कि इस न्यायालय में लंबित निम्नलिखित विवरण के अनुसार "
              "साक्षी का साक्ष्य अंकन जरिए वीडियो कांफ्रेंसिंग कराया जाना है:"), size=16, sa=3)

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    for f, v in [("ST NO / सत्र परीक्षण संख्या", st_no),
                 ("मुकदमा अपराध संख्या", mu_no),
                 ("राज्य बनाम", banam),
                 ("धारा", dhara),
                 ("जनपद", janpad),
                 ("तारीख पेशी", tareekh)]:
        i = [("ST NO / सत्र परीक्षण संख्या", st_no),
             ("मुकदमा अपराध संख्या", mu_no),
             ("राज्य बनाम", banam),
             ("धारा", dhara),
             ("जनपद", janpad),
             ("तारीख पेशी", tareekh)].index((f, v))
        set_cell(table.cell(i, 0), f)
        set_cell(table.cell(i, 1), v)

    fmt(doc, (f"\nअतः आपसे निवेदन है कि साक्षी का साक्ष्य अंकन {tareekh} को "
              f"{vc_sthan} जनपद {seva_janpad or '___'} में वीडियो कांफ्रेंसिंग के माध्यम से "
              f"करवाने हेतु आवश्यक निर्देश प्रदान करने की कृपा करें।"), size=16, sa=4)
    fmt(doc, f"भवदीय,\n{praapak}\n{today}",
        align=WD_PARAGRAPH_ALIGNMENT.RIGHT, size=16, sa=4)
    fmt(doc, "─" * 80, size=9, sa=1)
    fmt(doc, "प्रतिलिपि-", bold=True, size=15, sa=1)
    fmt(doc, (f"1. {prati_pad} जनपद {seva_janpad or '___'} को इस आशय के साथ प्रेषित कि "
              f"संबंधित अधीनस्थ को उपरोक्तानुसार निर्देश प्रदान करने का कष्ट करें।"), size=15, sa=1)
    fmt(doc, (f"2. {witness} दिनांक {tareekh} को समय {vc_samay} पर {vc_sthan} जनपद "
              f"{seva_janpad or '___'} में अपनी पहचान संबंधित दस्तावेज के साथ "
              f"वीडियो कांफ्रेंसिंग हेतु उपस्थित रहें।"), size=15, sa=2)
    fmt(doc, f"{praapak}\n{today}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, size=15)


# ── Main entry ───────────────────────────────────────────────────────────────

def generate_bulk(df_main: pd.DataFrame, janpad: str,
                  df_vc: pd.DataFrame = None) -> bytes:
    """
    df_main: main sheet DataFrame (columns per COL mapping)
    janpad: जनपद name
    df_vc: optional VC sheet DataFrame
    """
    # Validate columns
    missing = [c for c in REQUIRED_COLS if c not in df_main.columns]
    if missing:
        raise ValueError(f"इन columns की ज़रूरत है: {missing}")

    saman, bw, nbw, mrityu = [], [], [], []

    for _, row in df_main.iterrows():
        field = str(row.get(COL["saakshi"], "")).strip()
        if not field: continue
        stype = detect_type(field)
        # comma-separated witnesses → split
        names = [n.strip() for n in clean_prefix(field).split(',') if n.strip()]
        for name in names:
            r = row.copy()
            r[COL["saakshi"]] = name
            {"SAMAN": saman, "BW": bw, "NBW": nbw, "MRITYU": mrityu}[stype].append((r, stype))

    all_blocks = saman + bw + nbw + mrityu

    doc = base_doc()
    first = True

    for idx, (row, stype) in enumerate(all_blocks):
        if not first and idx % 2 == 0:
            doc.add_page_break()
        elif not first and idx % 2 == 1:
            p = doc.add_paragraph()
            r = p.add_run("─" * 110)
            r.font.size = Pt(7)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
        render_block(doc, row, stype, janpad)
        first = False

    # VC letters
    if df_vc is not None and not df_vc.empty:
        for _, row in df_vc.iterrows():
            doc.add_page_break()
            render_vc(doc, row, janpad)

    # Fix zoom
    zoom = doc.settings.element.find(qn('w:zoom'))
    if zoom is not None and zoom.get(qn('w:percent')) is None:
        zoom.set(qn('w:percent'), '100')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generate_from_excel(excel_bytes: bytes, janpad: str) -> bytes:
    """Convenience wrapper — reads Excel bytes directly."""
    xls = pd.ExcelFile(io.BytesIO(excel_bytes))
    df_main = pd.read_excel(xls, sheet_name=0).fillna("").rename(columns=clean_col)
    df_vc = None
    if "VC" in xls.sheet_names:
        df_vc = pd.read_excel(xls, sheet_name="VC").fillna("").rename(columns=clean_col)
    return generate_bulk(df_main, janpad, df_vc)
